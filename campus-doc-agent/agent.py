import os
import json
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple

import openai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

openai.api_key = os.getenv("OPENAI_API_KEY", "your-api-key-here")
LLM_MODEL = "gpt-4o-mini"

TEMPLATE_LIBRARY = {
    "实习报告": {
        "structure": [
            "标题: 实习报告",
            "一、实习基本信息",
            "   - 实习单位：{unit}",
            "   - 实习岗位：{position}",
            "   - 实习时间：{start_date} 至 {end_date}",
            "   - 指导教师：{advisor}",
            "二、实习目的",
            "三、实习内容",
            "四、实习总结与体会",
            "五、致谢",
            "格式要求：不少于800字，标题黑体三号，正文宋体小四，1.5倍行距"
        ],
        "required_fields": ["unit", "position", "start_date", "end_date", "advisor"],
        "min_words": 800
    },
    "保证书": {
        "structure": [
            "标题: 保证书",
            "致：{recipient}",
            "本人 {name}，学号 {student_id}，所在院系 {department}，就 {matter} 一事，作出如下保证：",
            "1. 本人深刻认识到该事件的严重性，并诚恳接受批评教育。",
            "2. 本人郑重承诺，今后将严格遵守学校各项规章制度，杜绝类似行为再次发生。",
            "3. 若再次违反，愿意接受学校相关纪律处分。",
            "特此保证。",
            "",
            "保证人（签名）：{name}",
            "日期：{date}"
        ],
        "required_fields": ["recipient", "name", "student_id", "department", "matter", "date"],
        "min_words": None,
        "tone": "严肃、诚恳，不得使用口语化表达"
    },
    "课程考核通知": {
        "structure": [
            "标题: 关于《{course_name}》课程考核安排的通知",
            "各位同学：",
            "现将《{course_name}》课程考核相关事宜通知如下：",
            "一、考核方式：{exam_type}",
            "二、考核时间：{exam_date} {exam_time}",
            "三、考核地点：{exam_location}",
            "四、注意事项：",
            "   1. 请携带学生证和必要文具；",
            "   2. 严格遵守考场纪律；",
            "   3. 其他说明：{notes}",
            "请同学们相互转告，按时参加考核。",
            "",
            "任课教师：{teacher}",
            "发布日期：{date}"
        ],
        "required_fields": ["course_name", "exam_type", "exam_date", "exam_time",
                            "exam_location", "notes", "teacher", "date"],
        "min_words": None,
        "info_completeness": "所有必填项必须完整，时间地点需明确，落款不能缺失"
    }
}

REQUIREMENT_EXTRACTION_PROMPT = """你是一个高校文书处理专家。请从以下用户输入中提取关键约束信息，并以JSON格式返回。
用户输入：
{user_input}

可选的文书类型：{doc_types}
请返回一个JSON对象，包含：
- doc_type: 文书类型
- extracted_fields: 提取的所有信息字段（字典）
- special_requirements: 用户额外格式或内容要求（字符串）

示例输出：
{{
  "doc_type": "保证书",
  "extracted_fields": {{
    "name": "张三",
    "student_id": "2021001",
    "department": "计算机学院",
    "matter": "旷课",
    "date": "2026-04-29"
  }},
  "special_requirements": "需要手写签名行"
}}
"""

VALIDATION_PROMPT = """你是严格的文书合规性审查专家。请审核以下文书。

文书类型：{doc_type}
文书内容：
{document}

具体要求：{requirements}

返回JSON：
- compliant: true/false
- issues: 问题列表
- suggestions: 修改建议列表
"""


def parse_json_safely(text: str) -> dict:
    match = re.search(r'```(?:json)?\s*(.*?)\s*```', text, re.DOTALL)
    if match:
        text = match.group(1)
    return json.loads(text)


def call_llm(prompt: str, max_tokens=1000, temperature=0.2) -> str:
    response = openai.ChatCompletion.create(
        model=LLM_MODEL,
        messages=[
            {"role": "system", "content": "你是一个专业的高校行政助手。"},
            {"role": "user", "content": prompt}
        ],
        max_tokens=max_tokens,
        temperature=temperature
    )
    return response.choices[0].message.content


def fill_template(doc_type: str, fields: dict) -> str:
    if doc_type not in TEMPLATE_LIBRARY:
        raise ValueError(f"不支持的文书类型: {doc_type}")
    lines = TEMPLATE_LIBRARY[doc_type]["structure"]
    filled = []
    for line in lines:
        try:
            filled.append(line.format_map(fields))
        except KeyError:
            filled.append(line)
    return "\n".join(filled)


def generate_word_document(content: str, filename: str) -> str:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    for para_text in content.split('\n'):
        if not para_text.strip():
            continue
        p = doc.add_paragraph()
        if para_text.strip().startswith("标题:"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(para_text.replace("标题:", "").strip())
            run.font.size = Pt(16)
            run.bold = True
        else:
            run = p.add_run(para_text)
            run.font.size = Pt(12)
    filepath = os.path.join(os.getcwd(), filename)
    doc.save(filepath)
    return filepath


def generate_markdown_file(content: str, filename: str) -> str:
    md_content = content.replace("标题:", "# ")
    filepath = os.path.join(os.getcwd(), filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(md_content)
    return filepath


class CampusDocAgent:
    def __init__(self):
        self.token_usage = 0
        self.doc_types = list(TEMPLATE_LIBRARY.keys())

    def process(self, user_input: str, user_info: Optional[dict] = None) -> Dict:
        result = {
            "status": "success",
            "doc_type": None,
            "document": None,
            "markdown_file": None,
            "word_file": None,
            "compliance_report": None,
            "final_suggestions": [],
            "token_used": 0
        }

        # Step 1: 需求解析
        extraction_prompt = REQUIREMENT_EXTRACTION_PROMPT.format(
            user_input=user_input,
            doc_types=", ".join(self.doc_types)
        )
        llm_output = call_llm(extraction_prompt, max_tokens=500)
        self._add_tokens(len(llm_output))

        try:
            extracted = parse_json_safely(llm_output)
        except json.JSONDecodeError:
            result["status"] = "error"
            result["final_suggestions"] = ["无法解析需求，请提供更明确的文书类型和信息。"]
            return result

        doc_type = extracted.get("doc_type")
        extracted_fields = extracted.get("extracted_fields", {})
        special_req = extracted.get("special_requirements", "")

        if doc_type not in TEMPLATE_LIBRARY:
            result["status"] = "error"
            result["final_suggestions"].append(f"暂不支持的文书类型: {doc_type}")
            return result

        if user_info:
            extracted_fields.update(user_info)

        # Step 2: 模板初稿
        try:
            draft = fill_template(doc_type, extracted_fields)
        except Exception as e:
            result["status"] = "error"
            result["final_suggestions"].append(f"模板填充失败: {str(e)}")
            return result

        # Step 3: LLM润色及内容扩展
        refine_prompt = f"""根据以下模板初稿和用户信息，生成一篇完整的文书。缺失占位符请合理补充。
文书类型：{doc_type}
模板初稿：
{draft}

用户额外信息：{json.dumps(extracted_fields, ensure_ascii=False)}
特殊要求：{special_req if special_req else "无"}
字数要求：{TEMPLATE_LIBRARY[doc_type].get('min_words', '无')}

请直接输出完整的文书内容，勿加解释。"""
        refined_content = call_llm(refine_prompt, max_tokens=1500)
        self._add_tokens(len(refined_content))
        result["document"] = refined_content

        # Step 4: 合规性校验
        validation_req = TEMPLATE_LIBRARY[doc_type].get("tone") or TEMPLATE_LIBRARY[doc_type].get("info_completeness", "")
        val_prompt = VALIDATION_PROMPT.format(
            doc_type=doc_type,
            document=refined_content,
            requirements=f"必填字段：{TEMPLATE_LIBRARY[doc_type]['required_fields']}。{validation_req}。特殊要求：{special_req}"
        )
        val_output = call_llm(val_prompt, max_tokens=500)
        self._add_tokens(len(val_output))

        try:
            compliance = parse_json_safely(val_output)
        except json.JSONDecodeError:
            compliance = {"compliant": True, "issues": [], "suggestions": []}

        result["compliance_report"] = compliance

        if not compliance.get("compliant", True):
            fix_prompt = f"""以下文书存在合规性问题，请根据建议修改，输出修改后的完整文书。
原始文书：
{refined_content}

问题列表：
{chr(10).join(compliance.get('issues', []))}
修改建议：
{chr(10).join(compliance.get('suggestions', []))}

直接输出修正后的文书。"""
            fixed_content = call_llm(fix_prompt, max_tokens=1500)
            self._add_tokens(len(fixed_content))
            result["document"] = fixed_content
            result["final_suggestions"] = compliance.get("suggestions", [])
        else:
            result["final_suggestions"] = ["文书已通过合规性校验。"]

        # Step 5: 生成文件
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        try:
            result["word_file"] = generate_word_document(result["document"], f"{doc_type}_{timestamp}.docx")
            result["markdown_file"] = generate_markdown_file(result["document"], f"{doc_type}_{timestamp}.md")
        except Exception as e:
            result["final_suggestions"].append(f"文件生成失败: {str(e)}")

        result["doc_type"] = doc_type
        result["token_used"] = self.token_usage
        return result

    def _add_tokens(self, text_length: int):
        self.token_usage += text_length // 2